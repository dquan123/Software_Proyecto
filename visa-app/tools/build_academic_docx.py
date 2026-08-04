from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "documento-academico.md"
OUTDIR = ROOT / "docs" / "entregables"
DOCX_PATH = OUTDIR / "Tarea_3_Pruebas_Automatizadas.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D3DF"
INK = "1F2937"
MUTED = "64748B"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Tarea 3 - Pruebas Automatizadas")
    set_font(run, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Universidad del Valle de Guatemala")
    set_font(run, size=14, bold=True, color=INK)

    for line in [
        "Facultad de Ingenieria",
        "Departamento de Ciencias de la Computacion",
        "CC3091 - Ingenieria de Software 2",
        "Semestre II - 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_font(run, size=12, color=INK)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Tarea 3\nPruebas Automatizadas")
    set_font(run, size=24, bold=True, color=BLUE)
    title.paragraph_format.space_before = Pt(28)
    title.paragraph_format.space_after = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Proyecto:\nApp de Apoyo y Orientacion para Solicitantes de Visa")
    set_font(run, size=14, bold=True, color=INK)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Integrantes:")
    set_font(run, size=12, bold=True, color=INK)

    for line in [
        "Diego Quan - 24336",
        "Diego Sebastian Guevara Casasola - 24128",
        "Juan Francisco Orozco Mijangos - 24647",
        "Norman Aguirre - 24479",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_font(run, size=11, color=INK)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ciudad de Guatemala, Guatemala\nJulio de 2026")
    set_font(run, size=12, color=INK)
    doc.add_page_break()


def add_toc(doc):
    heading = doc.add_heading("Contenido", level=1)
    heading.paragraph_format.keep_with_next = True
    sections = [
        "1. Resumen ejecutivo",
        "2. Introduccion",
        "3. Descripcion del proyecto",
        "4. Arquitectura real del sistema",
        "5. Importancia de las pruebas automatizadas",
        "6. Fundamentos de pruebas",
        "7. Herramientas investigadas para frontend",
        "8. Herramientas investigadas para backend",
        "9. Comparacion de herramientas",
        "10. Seleccion de herramientas",
        "11. Justificacion de Vitest y React Testing Library",
        "12. Justificacion de Jest y Supertest",
        "13. Configuracion implementada",
        "14. Estrategia de pruebas",
        "15. Pruebas implementadas en frontend",
        "16. Pruebas implementadas en backend",
        "17. Resultados",
        "18. Cobertura",
        "19. Ventajas",
        "20. Desventajas",
        "21. Tiempo empleado",
        "22. Hallazgos tecnicos",
        "23. Conclusiones",
        "24. Recomendaciones",
        "25. Referencias",
        "26. Anexos",
    ]
    for item in sections:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_font(run, size=10.5)
    doc.add_page_break()


def add_inline_runs(paragraph, text, bold=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=11, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=11, italic=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=11, bold=bold)


def add_markdown_table(doc, rows):
    data = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    if not data:
        return
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    set_table_width(table)
    for r_idx, row in enumerate(data):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(paragraph, value, bold=(r_idx == 0))
    doc.add_paragraph()


def add_code_block(doc, lines):
    if not lines:
        return
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_border(cell, color="E5E7EB")
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(lines))
    set_font(run, name="Consolas", size=8.5, color=INK)
    doc.add_paragraph()


def build():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_toc(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    table_rows = []
    skip_cover = False

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped == "## Portada":
            skip_cover = True
            continue
        if skip_cover:
            if stripped.startswith("## 1. "):
                skip_cover = False
            else:
                continue

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if re.match(r"^\|\s*-+", stripped):
                continue
            table_rows.append(stripped)
            continue
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

        if not stripped:
            continue

        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            continue
        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, number_match.group(2))
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped.replace("  ", " "))

    if table_rows:
        add_markdown_table(doc, table_rows)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
